# views.py (DataSearch)
# !/usr/bin/env python3
# coding=utf-8
# young.daniel@epa.gov
# py-lint: disable=C0301,E1101,R0901,W0613,W0622,C0411


"""Definition of views."""

from datetime import datetime
from docx import Document
from io import BytesIO
from openpyxl import Workbook
from os import path, remove
import tempfile
from wkhtmltopdf.views import PDFTemplateResponse
from xhtml2pdf import pisa
from zipfile import ZipFile
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.models import User
from django.http import HttpRequest, HttpResponseRedirect, HttpResponse, \
    JsonResponse
from django.shortcuts import render
from django.template.loader import get_template
from django.urls import reverse_lazy
from django.utils.decorators import method_decorator
from django.utils.text import slugify
from django.views.generic import TemplateView, ListView, CreateView, \
    DetailView, UpdateView, DeleteView
from constants.utils import download_file, download_files
from DataSearch.forms import ExistingDataForm
from DataSearch.models import ExistingData, ExistingDataSharingTeamMap, \
    Attachment, DataAttachmentMap
from DataSearch.settings import APP_DISCLAIMER
from qar5.models import SectionA
from teams.models import Team, TeamMembership

@login_required
@staff_member_required
def web_dev_tools(request, *args, **kwargs):
    """
    Method to redirect a user to a webpage.

    - Includes various custom admin functionality.
    - Includes button to remove extra new line characters/spaces from QAPP data.
    """
    return render(request, 'web_dev.html', {})


@login_required
@staff_member_required
def clean_qapps(request, *args, **kwargs):
    """
    Method to remove extra new line characters and spaces from QAPP data.

    - This has been fixed in the constants.
    - Strings that have already been inserted might need to be cleaned up.
    """
    # TODO: Fix all default data of all sections of all qapps
    sections_a = SectionA.objects.all()
    for sect in sections_a:
        # Clean Section A
        a3_clean = sect.a3.replace('\r\n', ' ').replace('    ', ' ')
        a9_clean = sect.a9.replace('\r\n', ' ').replace('    ', ' ')
        sect.a3 = a3_clean
        sect.a9 = a9_clean
        sect.save()
    return render(request, 'web_dev.html', {})


def get_existing_data_all():
    """Method to get all data regardless of user or team."""
    return ExistingData.objects.all()


def get_existing_data_user(user_id):
    """
    Method to get all Existing Data created by a User.
    """
    user = User.objects.get(id=user_id)
    return ExistingData.objects.filter(created_by=user)


def get_existing_data_team(team_id):
    """Method to get all data belonging to a team."""
    team = Team.objects.filter(id=team_id).first()
    if team:
        include_data = ExistingDataSharingTeamMap.objects.filter(
            team=team).values_list('data', flat=True)
        queryset = ExistingData.objects.filter(id__in=include_data)
        return queryset
    return []


def check_can_edit(data, user):
    """
    Method used to check if the provided user can
    edit the provided Data Search instance.

    All of the user's member teams are checked as well as the user's
    super user status or Data Search instance ownership status.
    """
    # Check if any of the user's teams have edit privilege:
    user_teams = TeamMembership.objects.filter(
        member=user).values_list('team', flat=True)

    for team in user_teams:
        data_team_map = ExistingDataSharingTeamMap.objects.filter(
            data=data, team=team).first()
        if data_team_map and data_team_map.can_edit:
            return True

    # Check if the user is super or owns the data:
    return user.is_superuser or data.created_by == user


class ExistingDataIndex(LoginRequiredMixin, TemplateView):
    """Class to return the first page of the Existing Data flow."""

    template_name = 'DataSearch/existing_data_index.html'

    def get_context_data(self, **kwargs):
        """
        Custom method override to send data to the template.

        - Specifically, want to send a list of users and teams to select from.
        """
        context = super().get_context_data(**kwargs)
        context['users'] = User.objects.all()
        context['teams'] = Team.objects.all()
        return context


class ExistingDataList(LoginRequiredMixin, ListView):
    """View for Existing Data Tracking Tool."""

    model = ExistingData
    context_object_name = 'existing_data_list'
    template_name = 'DataSearch/existing_data_list.html'

    def get_context_data(self, **kwargs):
        """
        Custom method override to send data to the template.

        - Specifically, want to send the user or team information for this list of data.
        """
        context = super().get_context_data(**kwargs)
        path = self.request.path.split('/')
        p_id = path[len(path) - 1]
        p_type = path[len(path) - 2]
        if p_type == 'user':
            context['p_user'] = User.objects.get(id=p_id)
        elif p_type == 'team':
            context['team'] = Team.objects.get(id=p_id)
        return context

    def get_queryset(self):
        """Add method docstring."""  # TODO add docstring.
        path = self.request.path.split('/')
        p_id = path[len(path) - 1]
        p_type = path[len(path) - 2]
        if p_type == 'user':
            return get_existing_data_user(p_id)
        if p_type == 'team':
            return get_existing_data_team(p_id)
        return get_existing_data_all()


class ExistingDataDetail(LoginRequiredMixin, DetailView):
    """View for viewing the details of a Existing data instance."""

    model = ExistingData
    template_name = 'DataSearch/existing_data_detail.html'

    def get_context_data(self, **kwargs):
        """
        Custom method override to send data to the template.

        - Specifically, want to send the user or team information for this list of data.
        """
        context = super().get_context_data(**kwargs)
        context['edit_disabled'] = ''
        if not check_can_edit(context['object'], self.request.user):
            context['edit_message'] = 'You cannot edit this data.'
            context['edit_disabled'] = 'disabled'

        referer = self.request.META['HTTP_REFERER'].split('/')
        p_id = referer[len(referer) - 1]
        p_type = referer[len(referer) - 2]
        if p_type == 'user':
            context['p_user'] = User.objects.get(id=p_id)
        elif p_type == 'team':
            context['team'] = Team.objects.get(id=p_id)
        return context


class ExistingDataEdit(LoginRequiredMixin, UpdateView):
    """View for editing the details of a Existing data instance."""

    model = ExistingData
    form_class = ExistingDataForm
    template_name = 'DataSearch/existing_data_edit.html'

    def get(self, request, *args, **kwargs):
        """
        Override default get request so we can verify the user has
        edit privileges, either through super status or team membership.
        """
        pk = kwargs.get('pk')
        data = ExistingData.objects.filter(id=pk).first()
        if check_can_edit(data, request.user):
            # We need to make sure we return a QappApproval form for editing:
            object = ExistingData.objects.filter(id=pk).first()
            return render(request, self.template_name,
                          {'object': object,
                           'form': ExistingDataForm(
                               instance=object, user=request.user)})

    @method_decorator(login_required)
    def post(self, request, *args, **kwargs):
        """Process the post request with a modified Existing Data form."""
        pk = kwargs.get('pk')
        instance = ExistingData.objects.filter(id=pk).first()
        if check_can_edit(instance, request.user):
            form = ExistingDataForm(request.POST, request.FILES, 
                                    instance=instance, user=request.user)
            if form.is_valid():
                obj = existing_form_process_teams_attachments(
                    request, form, instance)
                return HttpResponseRedirect('/existingdata/detail/%s' % obj.id)

            return render(request, self.template_name, {'form': form})

        reason = 'You cannot edit this data.'
        return HttpResponseRedirect('/existingdata/detail/%s' % pk, 401, reason)


def existing_form_process_teams_attachments(request, form, instance=None):
    """
    Custom method with shared code when saving an Existing Data instance.
    Assume form.is_valid() check passed outside this method.
    """
    obj = form.save(commit=False)
    obj.created_by = request.user
    obj.disclaimer_req = form.cleaned_data['disclaimer_req']
    obj.save()
    for field in request.FILES:
        file = request.FILES[field]
        # filename = fs.save(file.name, file).
        # uploaded_file_url = fs.url(filename).
        # Insert the attachment.
        attch = Attachment()
        attch.uploaded_by = request.user
        attch.name = file.name
        attch.file = file
        attch.save()
        # Insert DataAttachmentMap.
        attch_map = DataAttachmentMap()
        attch_map.data = obj
        attch_map.attachment = attch
        attch_map.save()

    # Prepare and insert teams data.
    teams = form.cleaned_data['teams']
    existing_teams = []
    # Get any existing teams mapped to this existing_data instance
    if instance:
        existing_teams = instance.teams.all()
        for team in existing_teams:
            if team not in teams:
                data_team_map = ExistingDataSharingTeamMap.objects.filter(
                    data=instance, team=team).first()
                if data_team_map:
                    data_team_map.delete()

    for team in teams:
        if not existing_teams or team not in existing_teams:
            data_team_map = ExistingDataSharingTeamMap()
            data_team_map.can_edit = True
            data_team_map.team = team
            data_team_map.data = obj
            data_team_map.save()

    return obj

class ExistingDataCreate(LoginRequiredMixin, CreateView):
    """Class for creating new Existing Data."""

    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        """Return a view with an empty form for creating a new Existing Data."""
        return render(request, "DataSearch/existing_data_create.html",
                      {'form': ExistingDataForm(user=request.user)})

    @method_decorator(login_required)
    def post(self, request, *args, **kwargs):
        """Process the post request with a new Existing Data form filled out."""
        form = ExistingDataForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            obj = existing_form_process_teams_attachments(request, form, None)
            return HttpResponseRedirect('/existingdata/detail/%s' % obj.id)
        return render(request, 'DataSearch/existing_data_create.html',
                      {'form': form})


class ExistingDataDelete(LoginRequiredMixin, DeleteView):
    """Class for deleting previously uploaded files."""

    model = ExistingData
    template_name = 'DataSearch/existing_data_confirm_delete.html'
    success_url = reverse_lazy('tracking_tool')
    # TODO: Check user permissions

    def get(self, request, *args, **kwargs):
        """
        Override the default get method so we can check if the user
        has permission to edit (delete) this object.
        """
        pk = kwargs.get('pk', None)
        if pk:
            object = ExistingData.objects.filter(id=pk).first()
            if object and check_can_edit(object, request.user):
                return render(request, self.template_name, {'object': object})
        return HttpResponseRedirect('/existingdata/detail/%s' % pk)


def home(request):
    """Renders the home page."""
    assert isinstance(request, HttpRequest)
    # ################################################
    # NOTE: REMOVE THIS AFTER DOCX EXPORT IS DONE!
    # ################################################
    # return HttpResponseRedirect('/qar5/exportdoc/1')
    return render(
        request,
        'index.html',
        {
            'title': 'Home Page',
            'year': datetime.now().year,
        }
    )


def contact(request):
    """Renders the contact page."""
    assert isinstance(request, HttpRequest)
    return render(
        request,
        'main/contact.html',
        {
            'title': 'Contact',
            'message': 'Your contact page.',
            'year': datetime.now().year,
        }
    )


def about(request):
    """Renders the about page."""
    assert isinstance(request, HttpRequest)
    return render(
        request,
        'main/about.html',
        {
            'title': 'About',
            'message': 'Your application description page.',
            'year': datetime.now().year,
        }
    )


def export_pdf(request, *args, **kwargs):
    """Function to export Existing  Data as a PDF document."""
    data_id = kwargs.get('pk', None)
    if data_id is None:
        # Export ALL ExistingData available for this user to PDF.
        data = get_existing_data_all()
        template = get_template('DataSearch/existing_data_pdf_multi.html')
        filename = 'export_existingdata_%s.pdf' % request.user.username
        # TODO get all attachment_ids
        attachment_ids = DataAttachmentMap.objects.all().values_list(
            'attachment', flat=True)
    else:
        data = ExistingData.objects.get(id=data_id)
        template = get_template('DataSearch/existing_data_pdf.html')
        filename = 'export_%s.pdf' % data.source_title
        # get attachment_ids related to just this data id
        attachment_ids = DataAttachmentMap.objects.filter(
            data_id=data_id).values_list('attachment', flat=True)

    context_dict = {'object': data}

    # NOTE: If returning attachments alongside the PDF, we will need to
    # redesign the PDF downloader. As it stands, the PDF is downloaded with
    # each response. We need a way to create the PDF before sending the
    # response, then we can combine all files (PDF and attachments), then
    # return them all at once.
    resp = PDFTemplateResponse(
        request=request,
        template=template,
        filename=filename,
        context=context_dict,
        show_content_in_browser=False,
        cmd_options={},
    )
    if not attachment_ids or len(attachment_ids) == 0:
        return resp

    # Else we need to create a PDF from template without sending response
    html = template.render(context_dict)
    result = BytesIO()
    content = BytesIO(html.encode('utf-8'))
    pdf = pisa.pisaDocument(content, result)
    if pdf.err:
        return resp

    # Create a zip archive to return multiple files: PDF, n attachments.
    zip_mem = BytesIO()
    archive = ZipFile(zip_mem, 'w')

    # Always add the generated PDF from above first:
    # Having trouble writing the PDF to archive due to bad encoding.
    # Possible solution is to manually write byte string to a temporary
    # file location, write that file to archive, then delete temp file.
    temp_file_name = "/tmp/temp_%s.pdf" % filename
    with open(temp_file_name, 'wb') as temp_file:
        temp_file.write(result.getvalue())

    archive.write(temp_file_name)

    # Delete the tempfile after creating/writing/zipping it.
    remove(temp_file_name)

    # Then add all attachments
    for a_id in attachment_ids:
        # Get the actual file from server file system
        attachment = Attachment.objects.get(id=a_id)
        # attachment.filefield.file
        try:
            file = attachment.file.file
            # Zip attachment:
            archive.write(file.name, path.basename(file.name))
            temp_do_something = True
        except FileNotFoundError:
            print('Attachment File Not Found!')

    archive.close()
    response = HttpResponse(zip_mem.getvalue(), content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename="%s.zip"' % filename
    response['Content-length'] = zip_mem.tell()
    return response


def export_excel(request, *args, **kwargs):
    """Function to export Existing Existing Data as an Excel sheet."""
    data_id = kwargs.get('pk', None)
    if data_id is None:
        # Export ALL ExistingData available for this user to Excel.
        data = get_existing_data_all()
        filename = 'export_existingdata_%s.xlsx' % request.user.username

    else:
        data = [ExistingData.objects.get(id=data_id)]
        filename = 'export_%s.xlsx' % data[0].source_title

    workbook = Workbook()
    sheet = workbook.active
    row = 1

    for datum in data:
        # Optionally add colors formatting before writing to cells.
        # Programmatically write results to the PDF.
        for name, value in datum.get_fields():
            sheet.cell(row=row, column=1).value = name
            sheet.cell(row=row, column=2).value = value
            row += 1

        if datum.disclaimer_req:
            sheet.cell(row=row, column=1).value = 'Disclaimer'
            repl_str = '\n                    '
            # Replace '\n                    ' with ' ' in the disclaimer
            sheet.cell(row=row, column=2).value = APP_DISCLAIMER.replace(repl_str, ' ')

        row += 1  # Add a blank space between each individual Data set

    # Now return the generated excel sheet to be downloaded.
    content_type = 'application/vnd.vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    response = HttpResponse(content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename="%s"' % filename
    sheet.title = filename.split('.')[0]
    workbook.save(response)
    return response


def export_docx(request, *args, **kwargs):
    """Function to export Existing Existing Data as a Word doc."""
    if 'user' in request.path:
        user_id = kwargs.get('pk', None)
        team_id = data_id = None
    elif 'team' in request.path:
        team_id = kwargs.get('pk', None)
        user_id = data_id = None
    else:
        data_id = kwargs.get('pk', None)
        team_id = user_id = None

    if data_id is None or user_id or team_id:
        if user_id:
            data_ids = get_existing_data_user(
                user_id).values_list('id', flat=True)
        else:
            data_ids = get_existing_data_team(
                team_id).values_list('id', flat=True)

        zip_mem = BytesIO()
        archive = ZipFile(zip_mem, 'w')
        for id in data_ids:
            resp = export_doc_single(request, pk=id)
            filename = resp['filename']
            if filename:
                temp_file_name = '%d_%s' % (id, filename)
                with tempfile.SpooledTemporaryFile() as tmp:
                    archive.writestr(temp_file_name, resp.content)

        archive.close()
        response = HttpResponse(
            zip_mem.getvalue(), content_type='application/force-download')
        response['Content-Disposition'] = \
            'attachment; filename="%s_datasearches.zip"' % request.user.username
        response['Content-length'] = zip_mem.tell()
        return response


# py-lint: disable=no-member
def export_doc_single(request, *args, **kwargs):
    """Function to export a single DataSearch object as a Word Docx file."""
    data_id = kwargs.get('pk', None)
    data = ExistingData.objects.filter(id=data_id).first()

    if not data:
        return HttpResponseRedirect(request)

    filename = '%s.docx' % slugify(data.source_title)

    document = Document()
    document.add_heading(
        'Existing Data Search: %s' % data.source_title, level=1)
    
    document.add_heading('User Work Office/Lab', level=3)
    document.add_paragraph(data.work)
    document.add_heading('Email Address', level=3)
    document.add_paragraph(data.email)
    document.add_heading('Phone Number', level=3)
    document.add_paragraph(data.phone)
    document.add_heading('Search Phrase', level=3)
    document.add_paragraph(data.search)
    document.add_heading('Source', level=3)
    document.add_paragraph(data.source.name)
    document.add_heading('Source Title', level=3)
    document.add_paragraph(data.source_title)
    document.add_heading('Keywords', level=3)
    document.add_paragraph(data.keywords)
    document.add_heading('URL', level=3)
    document.add_paragraph(data.url)
    document.add_heading('Citation', level=3)
    document.add_paragraph(data.citation)
    document.add_heading('Date Accessed', level=3)
    document.add_paragraph(str(data.date_accessed))
    document.add_heading('Comments', level=3)
    document.add_paragraph(data.comments)
    document.add_heading('Created by', level=3)
    document.add_paragraph(str(data.created_by))

    document.add_heading('Teams Shared With', level=3)
    for team in data.teams.all():
        document.add_paragraph(team.name)
        # List members?

    document.add_heading('Attachments', level=2)
    for attachment in data.attachments.all():
        document.add_paragraph(attachment.name)
        # List uploaded_by?
        # Make hyperlink to open included file?

    content_type = 'application/vnd.openxmlformats-officedocument.' + \
            'wordprocessingml.document'
    response = HttpResponse(content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % filename
    document.save(response)
    response['filename'] = filename
    return response


def attachments_download(request, *args, **kwargs):
    """
    Method for downloading attachments for a provided existing data search.
    - pk for existing data (required)
    - id for attachment (optional)
    """
    existing_id = kwargs.get('pk', None)
    attachment_id = kwargs.get('id', None)
    existing = ExistingData.objects.filter(id=existing_id).first()

    if existing:
        if attachment_id is None:
            # Export all uploads for this user:
            file_list = existing.attachments.all()
            zip_name = 'existingdata_%s_attachments' % existing_id
            return download_files(file_list, zip_name)

        else:
            # Export the upload specified, if it belongs to the user:
            file = existing.attachments.filter(id=attachment_id).first()
            if file:
                return download_file(file)

    return HttpResponse(request)


def attachment_delete(request, *args, **kwargs):
    """
    Method for deleting an attachment for a provided existing data search.
    - pk for existing data (required)
    - id for attachment (required)
    """
    existing_id = kwargs.get('pk', None)
    attachment_id = kwargs.get('id', None)
    existing = ExistingData.objects.filter(id=existing_id).first()

    if existing and check_can_edit(existing, request.user):
        file = existing.attachments.filter(id=attachment_id).first()
        if file:
            file.delete()
            existing.save()

    return HttpResponseRedirect('/existingdata/detail/%s' % existing_id)
