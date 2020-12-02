# views.py (DataSearch)
# !/usr/bin/env python3
# coding=utf-8
# young.daniel@epa.gov
# py-lint: disable=C0301,E1101,R0901,W0613,W0622,C0411


"""Definition of views."""

from docx import Document
from io import BytesIO
from openpyxl import Workbook
from os import path, remove
import tempfile
from wkhtmltopdf.views import PDFTemplateResponse
from xhtml2pdf import pisa
from zipfile import ZipFile
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.http import HttpResponseRedirect, HttpResponse
from django.template.loader import get_template
from django.utils.decorators import method_decorator
from django.utils.text import slugify
from DataSearch.models import ExistingData, Attachment, DataAttachmentMap
from DataSearch.settings import APP_DISCLAIMER
from DataSearch.views import get_existing_data_team, get_existing_data_user


def export_pdf_single(request, *args, **kwargs):
    """Function to export Existing  Data as a PDF document."""
    data_id = kwargs.get('pk', None)
    if data_id:
        data = ExistingData.objects.get(id=data_id)
        template = get_template('DataSearch/existing_data_pdf.html')

        filename = 'export_%s.pdf' % data.source_title
        # get attachment_ids related to just this data id
        attachment_ids = DataAttachmentMap.objects.filter(
            data_id=data_id).values_list('attachment', flat=True)

        context_dict = {'object': data}

        # NOTE: If returning attachments alongside the PDF, we will need to
        # redesign the PDF downloader. As it stands, the PDF is downloaded with
        # each response. We need a way to create the PDF before sending the
        # response, then we can combine all files (PDF and attachments), then
        # return them all at once.
        resp = PDFTemplateResponse(
            request=request,
            template=template,
            filename=filename,
            context=context_dict,
            show_content_in_browser=False,
            cmd_options={},
        )

        # Else we need to create a PDF from template without sending response
        html = template.render(context_dict)
        result = BytesIO()
        content = BytesIO(html.encode('utf-8'))
        pdf = pisa.pisaDocument(content, result)
        if pdf.err:
            return resp

        # Create a zip archive to return multiple files: PDF, n attachments.
        zip_mem = BytesIO()
        archive = ZipFile(zip_mem, 'w')

        # Always add the generated PDF from above first:
        archive.writestr(filename, result.getvalue())

        # Then add all attachments
        archive = add_attachments_to_zip(archive, attachment_ids)

        archive.close()
        response = HttpResponse(zip_mem.getvalue(),
                                content_type='application/force-download')
        response['Content-Disposition'] = 'attachment; filename="%s.zip"' % \
            filename
        response['Content-length'] = zip_mem.tell()
        response['filename'] = '%s.zip' % filename
        return response


def add_attachments_to_zip(archive, attachment_ids):
    """Add docstring."""  # TODO
    for a_id in attachment_ids:
        # Get the actual file from server file system
        attachment = Attachment.objects.get(id=a_id)
        # attachment.filefield.file
        try:
            file = attachment.file.file
            # Zip attachment:
            archive.write(file.name, path.basename(file.name))
        except FileNotFoundError:
            print('Attachment File Not Found!')
    return archive


def export_excel_single(request, *args, **kwargs):
    """Function to export Existing Existing Data as an Excel sheet."""
    data_id = kwargs.get('pk', None)
    if data_id:
        data = ExistingData.objects.get(id=data_id)
        filename = 'export_%s.xlsx' % data.source_title
        workbook = Workbook()
        sheet = workbook.active
        row = 1

        # Optionally add colors formatting before writing to cells.
        # Programmatically write results to the PDF.
        for name, value in data.get_fields():
            sheet.cell(row=row, column=1).value = name
            sheet.cell(row=row, column=2).value = value
            row += 1

        if data.disclaimer_req:
            sheet.cell(row=row, column=1).value = 'Disclaimer'
            repl_str = '\n                    '
            # Replace '\n                    ' with ' ' in the disclaimer
            sheet.cell(row=row, column=2).value = \
                APP_DISCLAIMER.replace(repl_str, ' ')

        # If no attachments, return the generated excel sheet.
        content_type = 'application/vnd.vnd.openxmlformats-' + \
                       'officedocument.spreadsheetml.sheet'
        response = HttpResponse(content_type=content_type)
        response['Content-Disposition'] = 'attachment; filename="%s"' % \
            filename
        sheet.title = slugify(filename.split('.')[0])
        workbook.save(response)
        response['filename'] = '%s.xlsx' % filename

        # get attachment_ids related to just this data id
        attachment_ids = DataAttachmentMap.objects.filter(
            data_id=data_id).values_list('attachment', flat=True)
        if attachment_ids:
            # Create a zip archive to return multiple files: PDF, n attachments
            zip_mem = BytesIO()
            archive = ZipFile(zip_mem, 'w')
            archive = add_attachments_to_zip(archive, attachment_ids)

            # Add excel to the zip
            # with tempfile.SpooledTemporaryFile() as tmp:
            archive.writestr(filename, response.content)

            archive.close()
            response = HttpResponse(zip_mem.getvalue(),
                                    content_type='application/force-download')
            response['Content-Disposition'] = \
                'attachment; filename="%s.zip"' % filename
            response['Content-length'] = zip_mem.tell()
            response['filename'] = '%s.zip' % filename
            return response

        return response


# py-lint: disable=no-member
def export_doc_single(request, *args, **kwargs):
    """Function to export a single DataSearch object as a Word Docx file."""
    data_id = kwargs.get('pk', None)
    data = ExistingData.objects.filter(id=data_id).first()

    if not data:
        return HttpResponseRedirect(request)

    filename = '%s.docx' % slugify(data.source_title)

    document = Document()
    styles = document.styles
    document.add_heading(
        'Existing Data Search: %s' % data.source_title, level=1)

    document.add_heading('User Work Office/Lab', level=3)
    document.add_paragraph(data.work)
    document.add_heading('Email Address', level=3)
    document.add_paragraph(data.email)
    document.add_heading('Phone Number', level=3)
    document.add_paragraph(data.phone)
    document.add_heading('Search Phrase', level=3)
    document.add_paragraph(data.search)
    document.add_heading('Source', level=3)
    document.add_paragraph(data.source.name)
    document.add_heading('Source Title', level=3)
    document.add_paragraph(data.source_title)
    document.add_heading('Keywords', level=3)
    document.add_paragraph(data.keywords)
    document.add_heading('URL', level=3)
    document.add_paragraph(data.url)
    document.add_heading('Citation', level=3)
    document.add_paragraph(data.citation)
    document.add_heading('Date Accessed', level=3)
    document.add_paragraph(str(data.date_accessed))
    document.add_heading('Comments', level=3)
    document.add_paragraph(data.comments)
    document.add_heading('Created by', level=3)
    document.add_paragraph(str(data.created_by))

    teams = data.teams.all()
    if teams:
        document.add_heading('Shared With Teams', level=3)
        table = document.add_table(rows=len(teams), cols=1)
        row = 0
        for team in teams:
            table.rows[row].cells[0].text = team.name
            # List members?
            row += 1

    # Create a zip archive to return multiple files: Docx, n attachments.
    zip_mem = BytesIO()
    archive = ZipFile(zip_mem, 'w')

    attachments = data.attachments.all()
    if attachments:
        document.add_heading('Attachments', level=2)
        table = document.add_table(rows=len(attachments), cols=1)
        row = 0
        for attachment in attachments:
            table.rows[row].cells[0].text = attachment.name
            # List uploaded_by?
            # Make hyperlink to open included file?
            row += 1
            try:
                file = attachment.file.file
                archive.write(file.name, path.basename(file.name))
            except FileNotFoundError:
                print('AttachmentFileNotFound')

    # Write the finalized docx to zip file:
    content_type = 'application/vnd.openxmlformats-officedocument.' + \
                   'wordprocessingml.document'
    fake_response = HttpResponse(content_type)
    fake_response['Content-Disposition'] = 'attachment; filename=%s' % filename
    document.save(fake_response)

    with tempfile.SpooledTemporaryFile() as tmp:
        archive.writestr(filename, fake_response.content)

    archive.close()

    response = HttpResponse(zip_mem.getvalue(),
                            content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename="%s.zip"' % \
        filename
    response['Content-length'] = zip_mem.tell()
    response['filename'] = '%s.zip' % filename
    return response


def export(request, *args, **kwargs):
    """
    Function to export multiple Existing Data as
    a zip file of the provided formats.
    """
    if 'user' in request.path:
        user_id = kwargs.get('pk', None)
        team_id = data_id = None
    elif 'team' in request.path:
        team_id = kwargs.get('pk', None)
        user_id = data_id = None
    else:
        data_id = kwargs.get('pk', None)
        team_id = user_id = None

    if data_id is None or user_id or team_id:
        if user_id:
            data_ids = get_existing_data_user(
                user_id).values_list('id', flat=True)
        else:
            data_ids = get_existing_data_team(
                team_id).values_list('id', flat=True)

        # TODO: Determine file type to set the method type
        if 'docx' in request.path:
            export_funct = export_doc_single
        elif 'pdf' in request.path:
            export_funct = export_pdf_single
        else:
            export_funct = export_excel_single

        zip_mem = BytesIO()
        archive = ZipFile(zip_mem, 'w')
        for id in data_ids:
            resp = export_funct(request, pk=id)
            filename = resp['filename']
            if filename:
                temp_file_name = '%d_%s' % (id, filename)
                with tempfile.SpooledTemporaryFile() as tmp:
                    archive.writestr(temp_file_name, resp.content)

        user = User.objects.filter(
            id=user_id).values('username').first()
        username = 'username'
        if user:
            username = user.get('username', 'username')

        archive.close()
        response = HttpResponse(
            zip_mem.getvalue(), content_type='application/force-download')
        response['Content-Disposition'] = \
            'attachment; filename="%s_datasearches.zip"' % username
        response['Content-length'] = zip_mem.tell()
        return response

    return HttpResponseRedirect(request)
